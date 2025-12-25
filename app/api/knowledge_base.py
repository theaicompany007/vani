"""Knowledge Base API routes for RAG integration"""
from flask import Blueprint, request, jsonify, current_app
import logging
import os
import tempfile
from pathlib import Path
import requests
from app.integrations.rag_client import get_rag_client
from app.auth import require_auth, require_use_case, get_current_user
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)

knowledge_base_bp = Blueprint('knowledge_base', __name__)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'doc', 'docx'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
        except ImportError:
            # Fallback to PyPDF2
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise


def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        # Try UTF-8 first, then fallback to other encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode text file with any supported encoding")
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        raise


def extract_text_from_doc(file_path):
    """Extract text from DOC/DOCX file"""
    try:
        import docx
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except ImportError:
        logger.error("python-docx library not installed. Please install it with: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"Error extracting text from DOC/DOCX: {e}")
        raise


def chunk_text(text, max_chunk_size=1000, overlap=200):
    """Split text into chunks with overlap"""
    if len(text) <= max_chunk_size:
        return [text]
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chunk_size
        chunk = text[start:end]
        
        # Try to break at sentence boundary
        if end < len(text):
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            break_point = max(last_period, last_newline)
            if break_point > max_chunk_size * 0.5:  # Only break if we're at least halfway
                chunk = text[start:start + break_point + 1]
                end = start + break_point + 1
        
        chunks.append(chunk.strip())
        start = end - overlap  # Overlap for context
    
    return chunks


@knowledge_base_bp.route('/api/knowledge-base/query', methods=['GET'])
@require_auth
def query_knowledge_base():
    """Query the knowledge base"""
    try:
        query = request.args.get('query', '')
        collections = request.args.getlist('collections') or None
        top_k = int(request.args.get('top_k', 5))
        
        if not query:
            return jsonify({'error': 'Query parameter is required', 'success': False}), 400
        
        # If no collections specified (i.e., "All Collections" selected), fetch ALL collections from RAG service
        # This ensures Google Drive collections are included in the search
        if not collections:
            rag_service_url = os.getenv('RAG_SERVICE_URL', 'https://rag.theaicompany.co')
            rag_api_key = os.getenv('RAG_API_KEY')
            
            if rag_api_key:
                try:
                    headers = {
                        'Content-Type': 'application/json',
                        'Authorization': f'Bearer {rag_api_key}',
                        'x-api-key': rag_api_key
                    }
                    # Try multiple endpoints to get ALL collections
                    endpoints_to_try = [
                        f"{rag_service_url}/rag/collections",
                        f"{rag_service_url}/collections",
                        f"{rag_service_url}/api/collections"
                    ]
                    
                    for endpoint in endpoints_to_try:
                        try:
                            response = requests.get(endpoint, headers=headers, timeout=10)
                            if response.status_code == 200:
                                data = response.json()
                                # Handle different response formats
                                all_collections = []
                                if isinstance(data, list):
                                    all_collections = data
                                elif isinstance(data, dict):
                                    # Handle different dict response formats
                                    logger.debug(f"Processing dict response with {len(data)} keys")
                                    logger.debug(f"Sample data: {list(data.items())[:2] if data else 'empty'}")
                                    
                                    if data:
                                        # Check if there's a 'collections' key with nested dict structure
                                        # Format: {'collections': {'case_studies': {'count': 5}, ...}, 'total_collections': 12}
                                        if 'collections' in data and isinstance(data['collections'], dict):
                                            collections_dict = data['collections']
                                            # Check if the nested dict has values that are dicts (with 'count' keys)
                                            first_nested_value = next(iter(collections_dict.values())) if collections_dict else None
                                            if first_nested_value and isinstance(first_nested_value, dict):
                                                # Extract collection names from the nested dict keys
                                                all_collections = list(collections_dict.keys())
                                                logger.info(f"✓ Extracted {len(all_collections)} collection names from nested 'collections' dict")
                                                logger.debug(f"Collections list: {all_collections}")
                                            else:
                                                # The 'collections' value might be a list
                                                all_collections = collections_dict if isinstance(collections_dict, list) else []
                                        else:
                                            # Check if top-level values are dicts (format: {'case_studies': {'count': 5}})
                                            try:
                                                first_value = next(iter(data.values()))
                                                logger.debug(f"First value type: {type(first_value)}, value: {str(first_value)[:200]}")
                                                
                                                if isinstance(first_value, dict):
                                                    # Check if nested value is also a dict (indicating {'collection': {'count': N}})
                                                    first_nested = next(iter(first_value.values())) if first_value else None
                                                    if first_nested and isinstance(first_nested, dict):
                                                        # This is the nested format, extract from top-level keys
                                                        all_collections = list(data.keys())
                                                        logger.info(f"✓ Extracted {len(all_collections)} collection names from top-level dict")
                                                    else:
                                                        # Top-level dict with simple values, extract keys
                                                        all_collections = list(data.keys())
                                                        logger.info(f"✓ Extracted {len(all_collections)} collection names from dict keys")
                                                    logger.debug(f"Collections list: {all_collections}")
                                                else:
                                                    # Try common keys for list of collections
                                                    all_collections = data.get('collections') or data.get('data') or data.get('result') or data.get('items') or []
                                                    if isinstance(all_collections, list):
                                                        logger.debug(f"Found collections list in common key: {len(all_collections)} items")
                                                    else:
                                                        logger.debug(f"Tried common keys, got: {type(all_collections)}")
                                            except StopIteration:
                                                logger.warning("Dict has no values")
                                                all_collections = []
                                            except Exception as e:
                                                logger.error(f"Error processing dict: {e}")
                                                all_collections = []
                                    else:
                                        all_collections = []
                                else:
                                    all_collections = []
                                
                                # Ensure we have a list, not a dict
                                if all_collections:
                                    if isinstance(all_collections, list):
                                        collections = all_collections
                                    elif isinstance(all_collections, dict):
                                        # If somehow still a dict, extract keys
                                        collections = list(all_collections.keys())
                                        logger.warning(f"Collections was still a dict, extracted keys: {collections}")
                                    else:
                                        logger.warning(f"Unexpected collections type: {type(all_collections)}")
                                        continue
                                    
                                    logger.info(f"Fetched {len(collections)} collections from RAG service for 'All Collections' query: {collections}")
                                    break
                                else:
                                    logger.warning(f"Failed to extract collections. Got: {type(all_collections)} - {all_collections}")
                        except Exception as e:
                            logger.debug(f"Failed to fetch collections from {endpoint}: {e}")
                            continue
                    
                    if not collections:
                        logger.warning("Could not fetch collections from RAG service, will let RAG client handle it")
                except Exception as e:
                    logger.warning(f"Error fetching all collections: {e}, will let RAG client handle it")
        
        # Final safety check: ensure collections is a list, not a dict
        if collections and isinstance(collections, dict):
            logger.warning(f"Collections is still a dict! Converting to list. Dict: {collections}")
            collections = list(collections.keys())
            logger.info(f"Converted to list: {collections}")
        elif collections and not isinstance(collections, list):
            logger.warning(f"Collections is not a list! Type: {type(collections)}, Value: {collections}")
            collections = None  # Let RAG client handle it
        
        rag_client = get_rag_client()
        # Pass None to let RAG client fetch all collections if we couldn't fetch them here
        logger.debug(f"Calling RAG client with collections (type: {type(collections)}): {collections}")
        result = rag_client.query(query, collections=collections, top_k=top_k)
        
        return jsonify({
            'success': True,
            'query': query,
            'results': result.get('results', {}),
            'total_results': result.get('total_results', 0)
        })
    except Exception as e:
        logger.error(f"Error querying knowledge base: {e}")
        return jsonify({'error': str(e), 'success': False}), 500


@knowledge_base_bp.route('/api/knowledge-base/collections', methods=['GET'])
@require_auth
def list_collections():
    """List available collections"""
    try:
        rag_service_url = os.getenv('RAG_SERVICE_URL', 'https://rag.theaicompany.co')
        rag_api_key = os.getenv('RAG_API_KEY')
        
        if not rag_api_key:
            return jsonify({
                'success': True,
                'collections': ['case_studies', 'services', 'company_profiles', 'industry_insights', 'faqs', 'platforms']
            })
        
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {rag_api_key}',
            'x-api-key': rag_api_key
        }
        
        # Try to get collections from RAG service
        # Try multiple possible endpoints
        endpoints_to_try = [
            f"{rag_service_url}/rag/collections",
            f"{rag_service_url}/collections",
            f"{rag_service_url}/api/collections"
        ]
        
        for endpoint in endpoints_to_try:
            try:
                logger.info(f"Trying to fetch collections from {endpoint}")
                response = requests.get(
                    endpoint,
                    headers=headers,
                    timeout=10
                )
                
                logger.info(f"RAG service response status from {endpoint}: {response.status_code}")
                
                if response.status_code == 200:
                    try:
                        data = response.json()
                        logger.info(f"RAG service response data: {data}")
                        
                        # Try different possible response formats
                        collections = None
                        if isinstance(data, list):
                            # Response is directly a list
                            collections = data
                        elif isinstance(data, dict):
                            # Check if it's a dict of collections with counts (like {'case_studies': {'count': 5}})
                            if all(isinstance(v, dict) for v in data.values()):
                                # Extract collection names from dict keys
                                collections = list(data.keys())
                            else:
                                # Try common keys for list of collections
                                collections = (data.get('collections') or 
                                            data.get('data') or 
                                            data.get('result') or
                                            data.get('items'))
                        
                        if collections and isinstance(collections, list) and len(collections) > 0:
                            logger.info(f"Successfully fetched {len(collections)} collections from RAG service: {collections}")
                            return jsonify({'success': True, 'collections': collections})
                        else:
                            logger.warning(f"RAG service returned empty or invalid collections from {endpoint}. Response: {data}, Extracted: {collections}")
                    except ValueError as json_error:
                        logger.error(f"Failed to parse JSON response from {endpoint}: {json_error}. Response text: {response.text[:500]}")
                elif response.status_code == 404:
                    # Endpoint doesn't exist, try next one
                    logger.debug(f"Endpoint {endpoint} returned 404, trying next endpoint")
                    continue
                else:
                    error_text = response.text[:500] if response.text else "No response body"
                    logger.warning(f"RAG service returned status {response.status_code} from {endpoint}: {error_text}")
            except requests.exceptions.Timeout:
                logger.warning(f"Timeout fetching collections from {endpoint}")
                continue
            except requests.exceptions.RequestException as e:
                logger.warning(f"Request error fetching collections from {endpoint}: {e}")
                continue
            except Exception as e:
                logger.error(f"Unexpected error fetching collections from {endpoint}: {e}", exc_info=True)
                continue
        
        # Fallback to default collections
        default_collections = ['case_studies', 'services', 'company_profiles', 'industry_insights', 'faqs', 'platforms']
        logger.info(f"Using fallback collections: {default_collections}")
        return jsonify({
            'success': True,
            'collections': default_collections
        })
    except Exception as e:
        logger.error(f"Error listing collections: {e}")
        return jsonify({'error': str(e), 'success': False}), 500


@knowledge_base_bp.route('/api/knowledge-base/upload', methods=['POST'])
@require_auth
def upload_documents():
    """Upload PDF/TXT files to knowledge base"""
    try:
        user = get_current_user()
        if not user:
            return jsonify({'error': 'Unauthorized', 'success': False}), 401
        
        # Get form data
        collection = request.form.get('collection', 'services')
        tags = request.form.get('tags', '')
        metadata_json = request.form.get('metadata', '{}')
        
        # Parse metadata
        import json
        try:
            metadata = json.loads(metadata_json) if metadata_json else {}
        except:
            metadata = {}
        
        # Add user info to metadata
        metadata['uploaded_by'] = getattr(user, 'email', 'unknown') if user else 'unknown'
        metadata['uploaded_at'] = str(Path().cwd())  # Will be replaced with timestamp
        
        # Parse tags
        tag_list = [tag.strip() for tag in tags.split(',') if tag.strip()] if tags else []
        
        # Get uploaded files
        if 'files[]' not in request.files:
            files = request.files.getlist('files')
        else:
            files = request.files.getlist('files[]')
        
        if not files or all(f.filename == '' for f in files):
            return jsonify({'error': 'No files provided', 'success': False}), 400
        
        results = []
        rag_service_url = os.getenv('RAG_SERVICE_URL', 'https://rag.theaicompany.co')
        rag_api_key = os.getenv('RAG_API_KEY')
        
        if not rag_api_key:
            return jsonify({'error': 'RAG service not configured', 'success': False}), 500
        
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {rag_api_key}',
            'x-api-key': rag_api_key
        }
        
        # Process each file
        for file in files:
            if file.filename == '':
                continue
            
            if not allowed_file(file.filename):
                results.append({
                    'filename': file.filename,
                    'success': False,
                    'error': 'File type not allowed. Only PDF and TXT files are supported.'
                })
                continue
            
            # Check file size
            file.seek(0, os.SEEK_END)
            file_size = file.tell()
            file.seek(0)
            
            if file_size > MAX_FILE_SIZE:
                results.append({
                    'filename': file.filename,
                    'success': False,
                    'error': f'File too large. Maximum size is {MAX_FILE_SIZE / 1024 / 1024}MB'
                })
                continue
            
            try:
                # Save file temporarily
                filename = secure_filename(file.filename)
                temp_dir = Path(tempfile.gettempdir())
                temp_file = temp_dir / f"vani_upload_{filename}"
                
                file.save(str(temp_file))
                
                # Extract text based on file type
                file_ext = filename.rsplit('.', 1)[1].lower()
                if file_ext == 'pdf':
                    text = extract_text_from_pdf(str(temp_file))
                elif file_ext in ['doc', 'docx']:
                    text = extract_text_from_doc(str(temp_file))
                else:  # txt
                    text = extract_text_from_txt(str(temp_file))
                
                if not text or not text.strip():
                    results.append({
                        'filename': filename,
                        'success': False,
                        'error': 'Could not extract text from file'
                    })
                    temp_file.unlink()
                    continue
                
                # Chunk the text
                chunks = chunk_text(text)
                
                # Prepare documents and metadatas for RAG service
                documents = []
                metadatas = []
                
                for i, chunk in enumerate(chunks):
                    documents.append(chunk)
                    chunk_metadata = {
                        **metadata,
                        'filename': filename,
                        'file_size': file_size,
                        'chunk_index': i,
                        'total_chunks': len(chunks),
                        'tags': tag_list,
                        'source': 'file_upload'
                    }
                    metadatas.append(chunk_metadata)
                
                # Send to RAG service
                payload = {
                    'collection': collection,
                    'documents': documents,
                    'metadatas': metadatas
                }
                
                response = requests.post(
                    f"{rag_service_url}/rag/add",
                    json=payload,
                    headers=headers,
                    timeout=60
                )
                
                if response.status_code == 200:
                    results.append({
                        'filename': filename,
                        'success': True,
                        'chunks_added': len(chunks),
                        'message': f'Successfully uploaded {len(chunks)} chunk(s)'
                    })
                else:
                    results.append({
                        'filename': filename,
                        'success': False,
                        'error': f'RAG service error: {response.status_code} - {response.text}'
                    })
                
                # Clean up temp file
                temp_file.unlink()
                
            except Exception as e:
                logger.error(f"Error processing file {file.filename}: {e}")
                results.append({
                    'filename': file.filename,
                    'success': False,
                    'error': str(e)
                })
        
        success_count = sum(1 for r in results if r.get('success'))
        return jsonify({
            'success': True,
            'results': results,
            'total_files': len(results),
            'successful': success_count,
            'failed': len(results) - success_count
        })
        
    except Exception as e:
        logger.error(f"Error uploading documents: {e}")
        return jsonify({'error': str(e), 'success': False}), 500


@knowledge_base_bp.route('/api/knowledge-base/ingest-url', methods=['POST'])
@require_auth
def ingest_url():
    """Scrape and ingest URL into knowledge base"""
    try:
        user = get_current_user()
        if not user:
            return jsonify({'error': 'Unauthorized', 'success': False}), 401
        
        data = request.get_json()
        url = data.get('url', '').strip()
        collection = data.get('collection', 'services')
        tags = data.get('tags', [])
        metadata = data.get('metadata', {})
        
        if not url:
            return jsonify({'error': 'URL is required', 'success': False}), 400
        
        # Validate URL
        if not url.startswith(('http://', 'https://')):
            return jsonify({'error': 'Invalid URL format', 'success': False}), 400
        
        rag_service_url = os.getenv('RAG_SERVICE_URL', 'https://rag.theaicompany.co')
        rag_api_key = os.getenv('RAG_API_KEY')
        
        if not rag_api_key:
            return jsonify({'error': 'RAG service not configured', 'success': False}), 500
        
        # Add user info to metadata
        metadata['ingested_by'] = getattr(user, 'email', 'unknown') if user else 'unknown'
        metadata['source'] = 'url_scraping'
        
        # Call RAG service ingest endpoint
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {rag_api_key}',
            'x-api-key': rag_api_key
        }
        
        payload = {
            'url': url,
            'collection': collection,
            'tags': tags if isinstance(tags, list) else [t.strip() for t in str(tags).split(',') if t.strip()],
            'metadata': metadata
        }
        
        response = requests.post(
            f"{rag_service_url}/ingest",
            json=payload,
            headers=headers,
            timeout=120  # URL scraping can take longer
        )
        
        if response.status_code == 200:
            result = response.json()
            return jsonify({
                'success': True,
                'url': url,
                'collection': collection,
                'message': 'URL successfully scraped and ingested',
                'result': result
            })
        else:
            return jsonify({
                'success': False,
                'error': f'RAG service error: {response.status_code} - {response.text}'
            }), 500
        
    except Exception as e:
        logger.error(f"Error ingesting URL: {e}")
        return jsonify({'error': str(e), 'success': False}), 500





